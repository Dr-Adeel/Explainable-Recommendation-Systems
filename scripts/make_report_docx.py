from docx import Document
from docx.shared import Pt

text = '''
**3.5 Outils et Technologies**

- **Langage & runtime:** Python 3.10+ (utilisé via venv dans le dépôt).
- **Science des données:** `pandas`, `numpy`, `scipy` pour manipulation tabulaire et matrices creuses (`CSR`).
- **Modélisation collaborative:** `implicit` (AlternatingLeastSquares).
- **Modélisation visuelle:** CLIP (hébergé via Hugging Face / modèles PyTorch) pour extraction d’embeddings d’images.
- **Indexation & recherche:** `scikit-learn` `NearestNeighbors` pour prototypage, possibilité d’évoluer vers `faiss` pour grande échelle.
- **API & service:** `FastAPI` + `uvicorn` pour exposer endpoints REST.
- **Interface:** prototype Streamlit (`frontend/app.py`) ; migration possible vers React + Vite (SPA).
- **Outils d’ingestion et scripts:** scripts Python sous `scripts/` (prétraitements, build CSR, entraînement ALS, évaluation).
- **Stockage des données:** format columnar `parquet` pour tables items/interactions, `npz` pour matrices CSR, `npy` pour facteurs de modèles.
- **Métriques & évaluation:** scripts utilitaires basés sur `numpy`/`pandas` pour calcul de Precision@K, Recall@K.
- **Logging & monitoring minimal:** logging Python (`logging`), endpoint `/health`.

**3.6 Réalisation Technique**

**3.6.1 Préparation des Interactions**

- Source: Amazon Fashion (subset réduit pour développement). Données brutes parsées et filtrées via `scripts/parse_amazon_raw.py` / `scripts/prepare_amazon_als_data.py`.
- Produits: fichiers sous `data/amazon/processed*` : `users.parquet`, `items.parquet`, `interactions_train.parquet`, `interactions_valid.parquet`, `interactions_test.parquet`.
- Nettoyage: conversion d’identifiants vers `user_idx`/`item_idx`, traitement des valeurs manquantes, conversion des ratings en score implicite (`value=1` par défaut ou rating->value si présent).
- Résumé dataset (extrait projet): users=5000, items=6441, train nnz=18266, valid nnz≈1062, test nnz≈831, items_with_images≈3938.

**3.6.2 Construction des Matrices CSR**

- Script: `scripts/build_amazon_csr.py`.
- Entrées: `interactions_train.parquet` + `interactions_meta.json` (nombre users/items).
- Processus: mapping des colonnes (`user_idx`,`item_idx`,`value`) → création d’une `csr_matrix((data,(rows,cols)), shape=(n_users,n_items))`.
- Sortie: `data/.../als/train_csr.npz`, plus `csr_stats.json` contenant shape, nnz, density et statistiques sur les valeurs.
- Validations: vérification d’absence de valeurs négatives (implicit ALS requiert non‑négativité), vérification d’alignement de shapes entre CSR et facteurs.

**Commandes (exemples):**

```
python scripts/build_amazon_csr.py --processed_dir data/amazon/processed_small --out_dir data/amazon/processed_small/als
```

**3.6.3 Entraînement ALS**

- Script: `scripts/train_amazon_als.py`.
- Approche: implicit ALS (confidence scaling) — convertir `X` en `Cui = (X * alpha).T` pour que `implicit` reçoive item-user matrix.
- Hyperparamètres courants (baseline): `factors=64`, `iterations=20`, `regularization=0.05`, `alpha=20.0`.
- Artefacts produits: `model/user_factors.npy`, `model/item_factors.npy`, `model/train_meta.json` (méta des hyperparams et shapes).
- Tests rapides: recommander pour `user_idx=0` et affichage de quelques recommandations de retour.

Commandes (exemples):

```
python scripts/train_amazon_als.py
```

**3.6.4 Indexation CLIP**

- Extraction embeddings: exécuter CLIP sur images produits → écriture `amazon_image_embeddings.parquet` (colonnes `item_idx`, `embedding`).
- Catalogue: enrichir embeddings avec métadonnées (title, category, image_path) dans `amazon_clip_catalog.parquet`.
- Index: construire `NearestNeighbors(metric='cosine')` ou `faiss.IndexFlatIP` après normalisation. Index entraîné et stocké, et matrices d’embeddings sauvegardées (`X` stacked in-memory).
- Intégration: à l’initialisation de l’API (`@app.on_event("startup")`) le module charge `amazon_image_embeddings.parquet` et `amazon_clip_catalog.parquet`, construit `_nn`, `_X`, `_item_ids` et mapping `_idx_by_item`.

**3.6.5 API FastAPI**

- Endpoints principaux:
  - `GET /health` — état & métadonnées (items_total, clip_ready, als_ready, shapes).
  - `GET /amazon/similar-items` — param: `item_idx,k,pool,filter_category` → renvoie top-k par similarité image.
  - `GET /amazon/recommend-user` — param: `user_idx,k,pool` → recommandations ALS top-k (filtre seen).
  - `GET /amazon/recommend-hybrid` — param: `item_idx,k,pool,alpha,beta,gamma,filter_category` → renvoie recommandations hybrides avec `parts` et `reason`.
- Implémentation: `src/api/amazon_api.py` a été étendu pour charger automatiquement artefacts ALS (cherche dans `data/.../als/model` et `data/.../processed_small/als/model`) et calculer popularity vector.
- Explicabilité: calcul des contributions normalisées et raison textuelle template via `_compose_reason()`.

Démarrage API (dev):

```
uvicorn src.api.amazon_api:app --host 127.0.0.1 --port 8001 --reload
```

**3.6.6 Interface Streamlit**

- Prototype: `frontend/app.py` fournit trois onglets: Image Similarity, User (ALS), Hybrid.
- Fonctionnalités: sélection d’un item d’exemple, override `item_idx` ou `user_idx`, réglage `k` et `pool`, ajustement des poids `alpha/beta/gamma` dans l’onglet Hybrid.
- Intégration: l’UI appelle les endpoints FastAPI via `requests`. Caching local pour `/health` et sample items via `st.cache_data`.
- Note: plan de migration vers React pour UI plus riche (déjà discuté).

Pour lancer:

```
streamlit run frontend/app.py
```

**3.6.7 Génération des Explications**

- Méthode simple utilisée: pour chaque candidat, calculer trois signaux (image, user/ALS, popularity), normaliser (min-max) chaque signal en [0,1], multiplier par poids (α/β/γ), convertir en parts relatives (somme → 1).
- Format retourné par l’API pour chaque reco:
  - `score`: valeur combinée brute,
  - `parts`: dict `{image:0.6, user:0.3, popularity:0.1}`,
  - `reason`: phrase template explicative (e.g., « Visually similar product (image similarity contributes 60%). »).
- Motivation: fournir une justification lisible et directement interprétable par l’utilisateur final ou par un UI simple (barres, pourcentages).
- Extensions futures: signatures d’importance fine (feature‑level), exemples de contre‑faits, ou intégration d’un explainability model (SHAP-like pour facteurs).

**3.7 Scénarios d’Utilisation**

**3.7.1 Similarité Image**

- Cas: un utilisateur veut visualiser produits visuellement proches d’un produit de référence.
- Flux: front sélectionne `item_idx` → `/amazon/similar-items` renvoie top-k visuels → UI affiche images + métadonnées.

**3.7.2 Recommandation Utilisateur**

- Cas: page profil utilisateur montrant recommandations personnalisées.
- Flux: backend charge `user_factors.npy` + `item_factors.npy`, pour `user_idx` calcule scores (dot product) en excluant items déjà vus → retourne top-k via `/amazon/recommend-user`.

**3.7.3 Recommandation Hybride**

- Cas: page produit « Vous aimerez aussi » mélangeant similarité visuelle et comportement collab.
- Flux: obtenir pool CLIP, calculer signaux, normaliser, combiner avec α/β/γ → renvoyer top-k avec `parts` et `reason`.

**3.7.4 Visualisation des Explications**

- UI: pour chaque vignette recommandée, afficher une mini‑barre montrant part de `image/user/popularity` et une phrase `reason`.
- Interaction: sliders pour ajuster α/β/γ et voir la liste se mettre à jour — utile pour démontrer sensibilité et pour tuning interactif.

**4. Résultats et Évaluation**

**4.1 Protocole d’Évaluation**

- Jeux: `train/valid/test` extraits des interactions (définis dans `data/amazon/processed*`).
- Metrices: Precision@K et Recall@K (définies dans `scripts/eval_amazon_als.py` — fonctions `precision_at_k`, `recall_at_k`).
- Recommandations: génération top‑k pour chaque utilisateur avec item exclusion (filtre seen).
- Baseline: Popularity baseline (top items by train frequency), comparée au modèle ALS.

Définitions:
- Precision@K = |relevant ∩ recommended_topK| / K
- Recall@K = |relevant ∩ recommended_topK| / |relevant|

**4.2 Résultats du Modèle ALS**

- Résumé chiffré (extraits d’exécution) :
  - Validation (valid): Precision@10 = 0.0089 ; Recall@10 = 0.0892
  - Test: Precision@10 = 0.0062 ; Recall@10 = 0.0617
- Interprétation: la précision est faible en valeur absolue — attendu sur datasets implicites et peu d’interactions par utilisateur — mais le rappel montre qu’une proportion non négligeable des items réellement pertinents est retrouvée dans le top‑10.

**4.3 Comparaison avec Popularity Baseline**

- Observations générales: la baseline par popularité affiche des performances inférieures (voir `scripts/eval_amazon_als.py` comparant precision/recall).
- Interprétation: ALS capture signaux collaboratifs au‑delà de la popularité brute, même si gains marginaux peuvent être modestes sur ce sous‑échantillon (dépend fortement de la densité du jeu d’entraînement et du seuil de pertinence).

**4.4 Analyse Qualitative**

- Exemples d’états: recommandations ALS peuvent proposer produits complémentaires non visuellement similaires (ex.: accessoires), CLIP propose items visuellement proches (couleurs, motifs).
- Scénarios où hybridation aide: quand un produit a des visuels ambigus (PLT) mais bénéficie d’un signal collaboratif fort, la combinaison hybride améliore pertinence perçue.
- Limitations observées: cold‑start pour items sans images ou nouveaux items sans interactions; embeddings CLIP sensibles à variations de fond et crop.

**4.5 Évaluation de l’Explicabilité**

- Mesure: qualitativement évaluée par cohérence des `reason` avec la recommandation (ex.: si `image` domine, la vignette est visuellement proche).
- Retours qualitatifs attendus: les explications templates aident à comprendre pourquoi une reco apparaît, mais restent simples; utilité élevée pour UX, limitée pour audits techniques.
- Propositions d’amélioration: collecte d’un petit panel d’utilisateurs pour évaluer utilité et confiance via questionnaire (Likert), mesurer si explications augmentent click‑through rate (CTR).

**4.5 Discussion**

- Points forts:
  - Architecture modulaire permettant d’itérer séparément sur CLIP, ALS et la couche de fusion.
  - Exposition via API standardisée (`FastAPI`) facilitant intégration front/back.
  - Explicabilité minimale directement intégrée au pipeline hybride.
- Limites:
  - Performances limitées par la densité du dataset et le volume d’interactions (sous‑échantillon).
  - Explications templates sont heuristiques et non causales.
  - Indexation CLIP actuelle (`NearestNeighbors`) suffisante pour prototypage, mais `faiss` souhaitable en production.
- Travaux futurs:
  - Tuning systematique d’α/β/γ via validation (grid search) et calibration des signaux.
  - Implémentation d’un explainability model (SHAP ou surrogate) et d’exemples utilisateurs pour meilleures justifications.
  - Passage à React pour UI, intégration A/B tests pour mesurer gains UX/CTR.
  - Déploiement conteneurisé (Docker) + orchestrateur si nécessaire, et pipeline CI pour rebuild d’artefacts.

'''

doc = Document()
for para in text.split('\n\n'):
    p = doc.add_paragraph()
    if para.strip().startswith('**') and para.strip().endswith('**'):
        # header
        run = p.add_run(para.strip().strip('*'))
        run.bold = True
        run.font.size = Pt(14)
    else:
        run = p.add_run(para)
        run.font.size = Pt(11)

out_path = 'reports/rapport_recommandation_hybride.docx'
import os
os.makedirs(os.path.dirname(out_path), exist_ok=True)
doc.save(out_path)
print(f'Written {out_path}')
